#!/usr/bin/env python3
"""
Convert SEO Brief Markdown to DOCX

This script converts the generated SEO brief markdown file to DOCX format
using the python-docx library.
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
import re

def convert_markdown_to_docx(markdown_file, docx_file=None):
    """Convert markdown file to DOCX format"""
    
    if not os.path.exists(markdown_file):
        print(f"❌ Markdown file not found: {markdown_file}")
        return False
    
    if not docx_file:
        docx_file = markdown_file.replace('.md', '.docx')
    
    print(f"📄 Converting {markdown_file} to {docx_file}...")
    
    # Read the markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a new Word document
    doc = Document()
    
    # Split content into lines
    lines = content.split('\n')
    
    # Process each line
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line - add paragraph break
            doc.add_paragraph()
            continue
        
        # Handle headers
        if line.startswith('# '):
            # H1 header
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('## '):
            # H2 header
            heading = doc.add_heading(line[3:], level=2)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('### '):
            # H3 header
            heading = doc.add_heading(line[4:], level=3)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('#### '):
            # H4 header
            heading = doc.add_heading(line[5:], level=4)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('##### '):
            # H5 header
            heading = doc.add_heading(line[6:], level=5)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('###### '):
            # H6 header
            heading = doc.add_heading(line[7:], level=6)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('---'):
            # Horizontal rule - add a line break
            doc.add_paragraph('─' * 50)
        elif line.startswith('- '):
            # Bullet point
            bullet_text = line[2:]
            # Handle bold text in bullet points
            bullet_text = process_formatting(bullet_text)
            doc.add_paragraph(bullet_text, style='List Bullet')
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            bold_text = line[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(bold_text)
            run.bold = True
        elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            # Italic text
            italic_text = line[1:-1]
            p = doc.add_paragraph()
            run = p.add_run(italic_text)
            run.italic = True
        else:
            # Regular paragraph
            if line:
                # Process formatting within the line
                processed_line = process_formatting(line)
                doc.add_paragraph(processed_line)
    
    # Save the document
    doc.save(docx_file)
    
    print(f"✅ Successfully converted to DOCX: {docx_file}")
    return True

def process_formatting(text):
    """Process markdown formatting within text"""
    
    # Handle bold text **text**
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    
    # Handle italic text *text*
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    
    # Handle code `text`
    text = re.sub(r'`(.*?)`', r'\1', text)
    
    # Handle links [text](url) - keep just the text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    return text

def main():
    """Main function"""
    print("📄 SEO Brief Markdown to DOCX Converter")
    print("=" * 50)
    
    # Look for the most recent brief file
    brief_files = [
        "gpu-vs-cpu-for-ai-brief.md",
        "gpu-vs-cpu-for-ai-live-sitemap-brief.md"
    ]
    
    markdown_file = None
    for file in brief_files:
        if os.path.exists(file):
            markdown_file = file
            break
    
    if not markdown_file:
        print("❌ No brief file found. Please run the brief generator first.")
        return
    
    print(f"📄 Found brief file: {markdown_file}")
    
    # Convert to DOCX
    success = convert_markdown_to_docx(markdown_file)
    
    if success:
        docx_file = markdown_file.replace('.md', '.docx')
        print(f"\n🎉 Conversion complete!")
        print(f"📄 DOCX file: {docx_file}")
        
        # Show file size
        if os.path.exists(docx_file):
            size = os.path.getsize(docx_file)
            print(f"📏 File size: {size:,} bytes")
    else:
        print("❌ Conversion failed")

if __name__ == "__main__":
    main()
