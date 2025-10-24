#!/usr/bin/env python3
"""
Enhanced DOCX Converter for SEO Briefs

This script converts markdown SEO briefs to well-formatted DOCX files
with proper styling and formatting.
"""

import os
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import re

def create_styled_document():
    """Create a document with custom styles"""
    doc = Document()
    
    # Define custom styles
    styles = doc.styles
    
    # Create a custom style for bullet points with better formatting
    try:
        bullet_style = styles.add_style('CustomBullet', WD_STYLE_TYPE.PARAGRAPH)
        bullet_style.font.name = 'Calibri'
        bullet_style.font.size = Pt(11)
        bullet_style.paragraph_format.left_indent = Inches(0.25)
    except:
        pass  # Style already exists
    
    return doc

def convert_markdown_to_enhanced_docx(markdown_file, docx_file=None):
    """Convert markdown file to enhanced DOCX format with better styling"""
    
    if not os.path.exists(markdown_file):
        print(f"❌ Markdown file not found: {markdown_file}")
        return False
    
    if not docx_file:
        docx_file = markdown_file.replace('.md', '.docx')
    
    print(f"📄 Converting {markdown_file} to DOCX: {docx_file}...")
    
    # Read the markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create a styled document
    doc = create_styled_document()
    
    # Split content into lines
    lines = content.split('\n')
    
    # Process each line
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line - add paragraph break
            doc.add_paragraph()
            continue
        
        # Handle headers
        if line.startswith('# '):
            # H1 header
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('## '):
            # H2 header
            heading = doc.add_heading(line[3:], level=2)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('### '):
            # H3 header
            heading = doc.add_heading(line[4:], level=3)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('#### '):
            # H4 header
            heading = doc.add_heading(line[5:], level=4)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('##### '):
            # H5 header
            heading = doc.add_heading(line[6:], level=5)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('###### '):
            # H6 header
            heading = doc.add_heading(line[7:], level=6)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        elif line.startswith('---'):
            # Horizontal rule - add a decorative line
            p = doc.add_paragraph()
            run = p.add_run('─' * 50)
            run.font.size = Pt(10)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        elif line.startswith('- '):
            # Bullet point with better formatting
            bullet_text = line[2:]
            bullet_text = process_enhanced_formatting(bullet_text)
            
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            p.style.font.size = Pt(11)
        elif line.startswith('**') and line.endswith('**'):
            # Bold text
            bold_text = line[2:-2]
            p = doc.add_paragraph()
            run = p.add_run(bold_text)
            run.bold = True
            run.font.size = Pt(11)
        elif line.startswith('*') and line.endswith('*') and not line.startswith('**'):
            # Italic text
            italic_text = line[1:-1]
            p = doc.add_paragraph()
            run = p.add_run(italic_text)
            run.italic = True
            run.font.size = Pt(11)
        else:
            # Regular paragraph
            if line:
                processed_line = process_enhanced_formatting(line)
                p = doc.add_paragraph(processed_line)
                p.style.font.size = Pt(11)
    
    # Save the document
    doc.save(docx_file)
    
    print(f"✅ Successfully converted to enhanced DOCX: {docx_file}")
    return True

def process_enhanced_formatting(text):
    """Process markdown formatting with enhanced styling"""
    
    # Handle bold text **text** - keep bold formatting
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    
    # Handle italic text *text* - keep italic formatting
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    
    # Handle code `text` - keep as regular text
    text = re.sub(r'`(.*?)`', r'\1', text)
    
    # Handle links [text](url) - keep just the text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    return text

def main():
    """Main function"""
    print("📄 SEO Brief DOCX Converter")
    print("=" * 50)
    
    # Look for the most recent brief file
    brief_files = [
        "gpu-vs-cpu-for-ai-brief.md"
    ]
    
    markdown_file = None
    for file in brief_files:
        if os.path.exists(file):
            markdown_file = file
            break
    
    if not markdown_file:
        print("❌ No brief file found. Please run the brief generator first.")
        return
    
    print(f"📄 Found brief file: {markdown_file}")
    
    # Convert to enhanced DOCX
    success = convert_markdown_to_enhanced_docx(markdown_file)
    
    if success:
        docx_file = markdown_file.replace('.md', '.docx')
        print(f"\n🎉 Conversion complete!")
        print(f"📄 DOCX file: {docx_file}")
        
        # Show file size
        if os.path.exists(docx_file):
            size = os.path.getsize(docx_file)
            print(f"📏 File size: {size:,} bytes")
    else:
        print("❌ Conversion failed")

if __name__ == "__main__":
    main()
