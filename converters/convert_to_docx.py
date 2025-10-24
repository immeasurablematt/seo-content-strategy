#!/usr/bin/env python3
"""
Quick DOCX Converter for SEO Briefs

Usage: python3 convert_to_docx.py [markdown_file]
"""

import os
import sys
from pathlib import Path
from docx import Document
import re

def convert_to_docx(markdown_file):
    """Convert markdown file to DOCX"""
    
    if not os.path.exists(markdown_file):
        print(f"❌ File not found: {markdown_file}")
        return False
    
    docx_file = markdown_file.replace('.md', '.docx')
    
    print(f"📄 Converting {markdown_file} to {docx_file}...")
    
    # Read markdown file
    with open(markdown_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create Word document
    doc = Document()
    
    # Process content
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        
        if not line:
            doc.add_paragraph()
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:], level=4)
        elif line.startswith('---'):
            doc.add_paragraph('─' * 50)
        elif line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            if line:
                # Clean up formatting
                clean_line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
                clean_line = re.sub(r'\*(.*?)\*', r'\1', clean_line)
                clean_line = re.sub(r'`(.*?)`', r'\1', clean_line)
                clean_line = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', clean_line)
                doc.add_paragraph(clean_line)
    
    # Save document
    doc.save(docx_file)
    
    print(f"✅ Converted to DOCX: {docx_file}")
    return True

def main():
    """Main function"""
    if len(sys.argv) > 1:
        markdown_file = sys.argv[1]
    else:
        # Default to the most recent brief file
        markdown_file = "gpu-vs-cpu-for-ai-brief.md"
    
    convert_to_docx(markdown_file)

if __name__ == "__main__":
    main()
